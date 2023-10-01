from docx import Document
from openpyxl import load_workbook

from docx.shared import Pt #Aumenta a letra
from docx.shared import RGBColor #Cor da letra

import win32com.client as win32 #Envio de email automatico
outlook = win32.Dispatch("outlook.application")


documento = "DadosAlunosEmail.xlsx" #Arquivo excel que vai ser utilizado
planilhaDadosAlunos = load_workbook(documento) #load workbook  is used when you have to access an MS Excel file and you want to open workbook for some operation.

aba_selecionada = planilhaDadosAlunos["Nomes"]



for linha in range(2, len(aba_selecionada["A"]) + 1): #Começa da linha 2 na coluna A, para pegar os 4 nomes precisamos ir até a linha 6 por isso o +1

    arquivoWord = Document("Certificado3.docx") #Abre o arquivo do word

    estilo = arquivoWord.styles["Normal"] #Seleciona o estilo

    nomeAluno = aba_selecionada['A%s' % linha].value #Pega o conteudo da celula e armazena na variavel
    dia = aba_selecionada['B%s' % linha].value
    mes = aba_selecionada['C%s' % linha].value
    ano = aba_selecionada['D%s' % linha].value
    curso = aba_selecionada['E%s' % linha].value
    professor = aba_selecionada['F%s' % linha].value
    emailAluno = aba_selecionada['G%s' % linha].value


    for paragrafo in arquivoWord.paragraphs:

        if "@nome" in paragrafo.text: #Verifica cada paragrafo do arquivo word em busca do @exemplo
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        primeiraParte = "Concluiu com sucesso o curso de "
        segundaParte = ", como carga horária de 20 horas, promovido pela escola de Cursos Online em "
        terceiraParte = f"{segundaParte} {dia} de {mes} de {ano}."

        if "@Text" in paragrafo.text:  # Verifica cada paragrafo do arquivo word em busca do @exemplo
            paragrafo.text = primeiraParte
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            cursoEditado = paragrafo.add_run(curso)
            cursoEditado.font.color.rgb = RGBColor(0, 112, 255)
            cursoEditado.underline = True
            cursoEditado.bold = True
            cursoEditado = paragrafo.add_run(terceiraParte)

        if "@Instrutor" in paragrafo.text:  # Verifica cada paragrafo do arquivo word em busca do @exemplo
            paragrafo.text = professor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminhoCertificados = "C:\\Users\\znaya\\Desktop\\Python\\Pycharm\\Automação WORD\\Certificados alterando 3 paragrafos\\" + nomeAluno + ".docx"
    arquivoWord.save(caminhoCertificados)



    emailOutlook = outlook.CreateItem(0) #Cria email
    emailOutlook.To = emailAluno #Pra quem
    emailOutlook.Subject = f"Aqui está seu certificado {nomeAluno}!" #Titulo do email
    emailOutlook.HTMLBody = f'''
        <p>Fala, {nomeAluno}! </p>
        <p>Segue seu certificado do curso de {curso}.</p>
        <p><b>Parábens</b> por ter participado!</p>
        '''

    #Adiciona anexo/certificado
    emailOutlook.Attachments.Add(caminhoCertificados)

    #salvando como rascunho #.send pra enviar
    emailOutlook.send


print("Certificados gerados, e-mails enviados")

